import gradio as gr
import pandas as pd
import sqlite3
import json
import re
import os
import unicodedata
import numpy as np
from contextlib import contextmanager
from docx import Document
from datetime import datetime
import easyocr
import cv2

# ==================== CONFIGURACIÓN ====================
DB_PATH = "patients.db"
JSON_SCHEMA_PATH = "sections.json"

# ==================== FUNCIONES DE BASE DE DATOS ====================
@contextmanager
def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
    finally:
        conn.close()

def init_db():
    with get_db_connection() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS patients (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                data TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_expediente ON patients(json_extract(data, '$.expediente'))")
        conn.execute("""
            CREATE TABLE IF NOT EXISTS clinical_notes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_id INTEGER NOT NULL,
                note_type TEXT,
                content TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (patient_id) REFERENCES patients(id)
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS clinical_snapshots (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_id INTEGER NOT NULL,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                glasgow INTEGER,
                rass INTEGER,
                cpot INTEGER,
                fc INTEGER,
                pam INTEGER,
                lactato REAL,
                FOREIGN KEY (patient_id) REFERENCES patients(id)
            )
        """)
        conn.commit()

def insert_patient(data):
    with get_db_connection() as conn:
        cursor = conn.execute(
            "INSERT INTO patients (data) VALUES (?) RETURNING id",
            (json.dumps(data),)
        )
        patient_id = cursor.fetchone()["id"]
        conn.commit()
        return patient_id

def update_patient(patient_id, data):
    with get_db_connection() as conn:
        conn.execute(
            "UPDATE patients SET data = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (json.dumps(data), patient_id)
        )
        conn.commit()

def get_patient(patient_id):
    with get_db_connection() as conn:
        row = conn.execute("SELECT id, data, created_at, updated_at FROM patients WHERE id = ?", (patient_id,)).fetchone()
        if row:
            return dict(row) | {"data": json.loads(row["data"])}
        return None

def check_expediente_exists(expediente, exclude_id=None):
    with get_db_connection() as conn:
        query = "SELECT 1 FROM patients WHERE json_extract(data, '$.expediente') = ?"
        params = [expediente]
        if exclude_id:
            query += " AND id != ?"
            params.append(exclude_id)
        row = conn.execute(query, params).fetchone()
        return row is not None

def add_note(patient_id, content, note_type="general"):
    with get_db_connection() as conn:
        conn.execute(
            "INSERT INTO clinical_notes (patient_id, note_type, content) VALUES (?, ?, ?)",
            (patient_id, note_type, content)
        )
        conn.commit()
        return True

# ==================== CARGA DEL ESQUEMA JSON ====================
def load_json_schema():
    if os.path.exists(JSON_SCHEMA_PATH):
        with open(JSON_SCHEMA_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    else:
        # Esquema mínimo (para prueba)
        return {
            "sections": [
                {"name": "Datos Contextuales", "fields": [
                    {"name": "nombre_completo", "label": "Nombre Completo"},
                    {"name": "fecha_nacimiento", "label": "Fecha Nacimiento"},
                    {"name": "edad", "label": "Edad"},
                    {"name": "sexo", "label": "Sexo"},
                    {"name": "curp", "label": "CURP"},
                    {"name": "expediente", "label": "Expediente"}
                ]},
                {"name": "Hídrico y Renal", "fields": [
                    {"name": "sodio", "label": "Sodio"},
                    {"name": "potasio", "label": "Potasio"},
                    {"name": "cloro", "label": "Cloro"},
                    {"name": "creatinina", "label": "Creatinina"}
                ]},
                {"name": "Hematológico e Infeccioso", "fields": [
                    {"name": "pcr", "label": "PCR"},
                    {"name": "pct", "label": "Procalcitonina"},
                    {"name": "leucocitos", "label": "Leucocitos"},
                    {"name": "hemoglobina", "label": "Hemoglobina"},
                    {"name": "plaquetas", "label": "Plaquetas"}
                ]},
                {"name": "Evaluación inicial", "fields": [
                    {"name": "diagnostico_ingreso", "label": "Diagnóstico"},
                    {"name": "sofa_ingreso", "label": "SOFA"},
                    {"name": "sofa2_ingreso", "label": "SOFA II"},
                    {"name": "apache2_ingreso", "label": "APACHE II"}
                ]},
                {"name": "Datos de Egreso", "stage": "discharge", "fields": [
                    {"name": "fecha_egreso_uci", "label": "Fecha egreso UCI"},
                    {"name": "diagnostico_egreso", "label": "Diagnóstico egreso"},
                    {"name": "sodio_egreso", "label": "Sodio egreso"}
                ]}
            ]
        }

schema = load_json_schema()
ingreso_field_names = []
egreso_field_names = []
for sec in schema["sections"]:
    for f in sec.get("fields", []):
        field_stage = f.get("stage", sec.get("stage", "both"))
        if field_stage in ("admission", "both"):
            ingreso_field_names.append(f["name"])
        if field_stage in ("discharge", "both"):
            egreso_field_names.append(f["name"])

# ==================== FUNCIONES DE NORMALIZACIÓN ====================
def normalize_text(text):
    """Elimina acentos, convierte a mayúsculas y quita puntos innecesarios."""
    if not isinstance(text, str):
        return text
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('utf-8')
    text = text.upper()
    text = re.sub(r'\.(?=\s|$)', '', text)
    text = re.sub(r'[^\w\s\-]', '', text)
    return text.strip()

# ==================== EXTRACCIÓN DE DATOS DESDE DOCX (CENSO) ====================
def parse_nombre_cell(text):
    data = {}
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line and not re.search(r'AÑOS|CURP|FN:|EXP:', line.upper()):
            data['nombre_completo'] = normalize_text(line)
            break
    edad_match = re.search(r'(\d+)\s*AÑOS', text, re.IGNORECASE)
    if edad_match:
        data['edad'] = int(edad_match.group(1))
    curp_match = re.search(r'CURP:\s*([A-Z0-9]{18})', text, re.IGNORECASE)
    if curp_match:
        data['curp'] = curp_match.group(1).upper()
    fn_match = re.search(r'FN:\s*(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE)
    if fn_match:
        data['fecha_nacimiento'] = fn_match.group(1)
    exp_match = re.search(r'EXP:\s*(\d+)', text, re.IGNORECASE)
    if exp_match:
        data['expediente'] = exp_match.group(1)
    return data

def extract_lab_values(text):
    data = {}
    patterns = {
        'sodio': r'NA\s*[:]?\s*(\d+(?:\.\d+)?)',
        'potasio': r'K\s*[:]?\s*(\d+(?:\.\d+)?)',
        'cloro': r'CL\s*[:]?\s*(\d+(?:\.\d+)?)',
        'creatinina': r'CR\s*[:]?\s*(\d+(?:\.\d+)?)',
        'pcr': r'PCR\s*[:]?\s*(\d+(?:\.\d+)?)',
        'pct': r'PROCALCITONINA\s*[:]?\s*(\d+(?:\.\d+)?)',
        'leucocitos': r'LEUCOCITOS\s*[:]?\s*(\d+(?:\.\d+)?)',
        'hemoglobina': r'HB\s*[:]?\s*(\d+(?:\.\d+)?)',
        'plaquetas': r'PLAQUETAS\s*[:]?\s*(\d+(?:\.\d+)?)',
        'bun': r'BUN\s*[:]?\s*(\d+(?:\.\d+)?)',
        'urea': r'UREA\s*[:]?\s*(\d+(?:\.\d+)?)',
    }
    for key, pat in patterns.items():
        match = re.search(pat, text, re.IGNORECASE)
        if match:
            try:
                data[key] = float(match.group(1))
            except:
                data[key] = match.group(1)
    return data

def extract_scale_values(text):
    data = {}
    patterns = {
        'sofa_ingreso': r'SOFA\s*[:]?\s*(\d+)',
        'sofa2_ingreso': r'SOFA\s*2\s*[:]?\s*(\d+)',
        'apache2_ingreso': r'APACHE\s*II\s*[:]?\s*(\d+)',
    }
    for key, pat in patterns.items():
        match = re.search(pat, text, re.IGNORECASE)
        if match:
            data[key] = int(match.group(1))
    return data

def extract_diagnosis(text):
    lines = text.split('\n')
    diag_lines = []
    for line in lines:
        line = line.strip()
        if line and re.match(r'^[\-\d\.]+\s*', line):
            diag_lines.append(line)
    if diag_lines:
        return ' '.join(diag_lines)
    return text.strip()

def process_docx(file_path):
    doc = Document(file_path)
    if not doc.tables:
        return []
    table = doc.tables[0]
    header_cells = table.rows[0].cells
    headers = [cell.text.strip() for cell in header_cells]
    col_indices = {}
    for i, h in enumerate(headers):
        h_upper = h.upper()
        if 'CAMA' in h_upper:
            col_indices['cama'] = i
        elif 'NOMBRE' in h_upper:
            col_indices['nombre'] = i
        elif 'DIAGNÓSTICO' in h_upper or 'DIAGNOSTICO' in h_upper:
            col_indices['diagnostico'] = i
        elif 'FI: HOSP' in h_upper:
            col_indices['fi_hosp'] = i
        elif 'FI: UCI' in h_upper:
            col_indices['fi_uci'] = i
        elif 'ESTADO DE SALUD' in h_upper:
            col_indices['estado'] = i
        elif 'PRONÓSTICO' in h_upper:
            col_indices['pronostico'] = i
        elif 'AMV' in h_upper:
            col_indices['amv'] = i
        elif 'ANTIBIÓTICOS' in h_upper or 'CULTIVOS' in h_upper:
            col_indices['antibioticos'] = i
        elif 'PENDIENTES' in h_upper:
            col_indices['pendientes'] = i
        elif 'LABORATORIOS' in h_upper:
            col_indices['laboratorios'] = i
    patients = []
    for row in table.rows[1:]:
        cells = row.cells
        patient = {}
        if 'cama' in col_indices:
            patient['cama'] = cells[col_indices['cama']].text.strip()
        if 'nombre' in col_indices:
            nombre_text = cells[col_indices['nombre']].text.strip()
            patient.update(parse_nombre_cell(nombre_text))
        if 'diagnostico' in col_indices:
            diag_text = cells[col_indices['diagnostico']].text.strip()
            patient['diagnostico_ingreso'] = extract_diagnosis(diag_text)
        if 'fi_hosp' in col_indices:
            patient['fecha_ingreso_hosp'] = cells[col_indices['fi_hosp']].text.strip()
        if 'fi_uci' in col_indices:
            patient['fecha_ingreso'] = cells[col_indices['fi_uci']].text.strip()
        if 'estado' in col_indices:
            patient['estado_salud'] = cells[col_indices['estado']].text.strip()
        if 'pronostico' in col_indices:
            pronostico_text = cells[col_indices['pronostico']].text.strip()
            patient.update(extract_scale_values(pronostico_text))
        if 'amv' in col_indices:
            patient['amv'] = cells[col_indices['amv']].text.strip()
        if 'antibioticos' in col_indices:
            patient['antibioticos'] = cells[col_indices['antibioticos']].text.strip()
        if 'pendientes' in col_indices:
            patient['pendientes'] = cells[col_indices['pendientes']].text.strip()
        if 'laboratorios' in col_indices:
            lab_text = cells[col_indices['laboratorios']].text.strip()
            patient.update(extract_lab_values(lab_text))
            patient['laboratorios_text'] = lab_text
        patients.append(patient)
    return patients

def extract_date_from_filename(filename):
    base = os.path.basename(filename)
    match = re.search(r'(\d{2})[.\-](\d{2})[.\-](\d{4})', base)
    if match:
        day, month, year = match.groups()
        try:
            return datetime(int(year), int(month), int(day))
        except:
            return datetime.max
    return datetime.max

def merge_patient_records(records):
    if not records:
        return {}
    base = records[0].copy()
    for record in records[1:]:
        for key, value in record.items():
            if key not in base or base[key] is None or base[key] == "" or (isinstance(base[key], float) and np.isnan(base[key])):
                if value is not None and value != "" and not (isinstance(value, float) and np.isnan(value)):
                    base[key] = value
    return base

def process_multiple_docx(files):
    all_patients_with_date = []
    for file in files:
        if file.name.endswith('.docx'):
            date = extract_date_from_filename(file.name)
            patients = process_docx(file.name)
            for p in patients:
                all_patients_with_date.append((date, p))
    total_records = len(all_patients_with_date)
    all_patients_with_date.sort(key=lambda x: x[0])
    groups = {}
    for date, patient in all_patients_with_date:
        pid = patient.get('expediente') or patient.get('curp')
        if not pid:
            nombre = patient.get('nombre_completo', '')
            fn = patient.get('fecha_nacimiento', '')
            pid = f"{nombre}_{fn}" if nombre or fn else f"unknown_{len(groups)}"
        if pid not in groups:
            groups[pid] = []
        groups[pid].append(patient)
    merged_patients = []
    for pid, records in groups.items():
        merged = merge_patient_records(records)
        merged_patients.append(merged)
    return merged_patients, total_records

# ==================== FUNCIONES PARA CENSO (DATAFRAME Y GUARDADO) ====================
def create_dataframe_from_patients(patients):
    columns = list(set(ingreso_field_names + ['cama', 'fecha_ingreso', 'fecha_ingreso_hosp', 'estado_salud', 'amv', 'antibioticos', 'pendientes', 'laboratorios_text', 'notas']))
    priority = ['nombre_completo', 'edad', 'sexo', 'curp', 'expediente', 'fecha_nacimiento',
                'diagnostico_ingreso', 'sodio', 'potasio', 'cloro', 'creatinina',
                'pcr', 'pct', 'leucocitos', 'hemoglobina', 'plaquetas',
                'sofa_ingreso', 'sofa2_ingreso', 'apache2_ingreso',
                'cama', 'fecha_ingreso', 'fecha_ingreso_hosp', 'estado_salud', 'amv',
                'antibioticos', 'pendientes', 'laboratorios_text', 'notas']
    final_columns = [c for c in priority if c in columns]
    df = pd.DataFrame(patients)
    for col in final_columns:
        if col not in df.columns:
            df[col] = ""
    df = df.fillna("")
    df = df[final_columns]
    return df

def save_to_database_from_df(df):
    resultados = []
    for idx, row in df.iterrows():
        data = {}
        for col in df.columns:
            if col != 'notas' and col != 'laboratorios_text' and pd.notna(row[col]) and row[col] != "":
                val = row[col]
                if isinstance(val, str):
                    if re.match(r'^\d+$', val):
                        val = int(val)
                    elif re.match(r'^\d+\.\d+$', val):
                        val = float(val)
                data[col] = val
        if 'expediente' not in data or not data['expediente']:
            resultados.append(f"Fila {idx+1}: Falta expediente, no se guarda")
            continue
        existe = check_expediente_exists(data['expediente'])
        if existe:
            with get_db_connection() as conn:
                row_db = conn.execute(
                    "SELECT id FROM patients WHERE json_extract(data, '$.expediente') = ?",
                    (data['expediente'],)
                ).fetchone()
            if row_db:
                patient_id = row_db['id']
                current = get_patient(patient_id)
                if current:
                    current_data = current["data"]
                    current_data.update(data)
                    update_patient(patient_id, current_data)
                    msg = f"Fila {idx+1}: Actualizado paciente {data.get('nombre_completo', '')} (ID: {patient_id})"
                else:
                    msg = f"Fila {idx+1}: Error al recuperar paciente"
            else:
                msg = f"Fila {idx+1}: Expediente existe pero no se pudo obtener ID"
        else:
            new_id = insert_patient(data)
            msg = f"Fila {idx+1}: Insertado nuevo paciente {data.get('nombre_completo', '')} (ID: {new_id})"
            patient_id = new_id
        if pd.notna(row['notas']) and row['notas'].strip():
            add_note(patient_id, row['notas'].strip(), note_type='entrega')
            msg += " + nota guardada"
        resultados.append(msg)
    return "\n".join(resultados)

# ==================== FUNCIONES PARA PROCESAR EGRESO ====================
def extract_text_from_image(image_path):
    reader = easyocr.Reader(['es'])
    img = cv2.imread(image_path)
    result = reader.readtext(img, detail=0, paragraph=True)
    return "\n".join(result)

def identify_patient_from_text(text, conn):
    text_norm = normalize_text(text)
    exp_match = re.search(r'EXP[:\s]*(\d+)', text, re.IGNORECASE)
    if exp_match:
        expediente = exp_match.group(1)
        row = conn.execute("SELECT id, data FROM patients WHERE json_extract(data, '$.expediente') = ?", (expediente,)).fetchone()
        if row:
            return row['id'], json.loads(row['data'])
    curp_match = re.search(r'CURP[:\s]*([A-Z0-9]{18})', text, re.IGNORECASE)
    if curp_match:
        curp = curp_match.group(1)
        row = conn.execute("SELECT id, data FROM patients WHERE json_extract(data, '$.curp') = ?", (curp,)).fetchone()
        if row:
            return row['id'], json.loads(row['data'])
    lines = text.split('\n')
    for line in lines[:3]:
        words = line.strip().split()
        if len(words) >= 2:
            nombre_candidato = ' '.join(words[:2])
            rows = conn.execute("SELECT id, data FROM patients WHERE json_extract(data, '$.nombre_completo') LIKE ?", (f'%{nombre_candidato}%',)).fetchall()
            if rows:
                return rows[0]['id'], json.loads(rows[0]['data'])
    return None, None

def extract_discharge_data(text):
    data = {}
    patterns = {
        'fecha_egreso_uci': r'fecha de egreso[:\s]*(\d{2}/\d{2}/\d{4})',
        'diagnostico_egreso': r'diagn[oó]stico de egreso[:\s]*(.+?)(?=\n|$)',
        'plan_egreso': r'plan de egreso[:\s]*(.+?)(?=\n|$)',
        'condicion_egreso': r'condici[oó]n (?:de )?egreso[:\s]*(\w+)',
        'destino_egreso': r'destino[:\s]*(\w+)',
        'temperatura_egreso': r'temperatura[:\s]*(\d+(?:\.\d+)?)',
        'fc_egreso': r'fc[:\s]*(\d+)',
        'fr_egreso': r'fr[:\s]*(\d+)',
        'tas_egreso': r'tas[:\s]*(\d+)',
        'tad_egreso': r'tad[:\s]*(\d+)',
        'sao2_egreso': r'sao2[:\s]*(\d+)',
        'fio2_egreso': r'fio2[:\s]*(\d+)',
        'hemoglobina_egreso': r'hemoglobina[:\s]*(\d+(?:\.\d+)?)',
        'hematocrito_egreso': r'hematocrito[:\s]*(\d+(?:\.\d+)?)',
        'leucocitos_egreso': r'leucocitos[:\s]*(\d+(?:\.\d+)?)',
        'plaquetas_egreso': r'plaquetas[:\s]*(\d+(?:\.\d+)?)',
        'neutrofilos_egreso': r'neutr[oó]filos[:\s]*(\d+(?:\.\d+)?)',
        'linfocitos_egreso': r'linfocitos[:\s]*(\d+(?:\.\d+)?)',
        'pcr_egreso': r'pcr[:\s]*(\d+(?:\.\d+)?)',
        'pct_egreso': r'procalcitonina[:\s]*(\d+(?:\.\d+)?)',
        'sodio_egreso': r'sodio[:\s]*(\d+(?:\.\d+)?)',
        'potasio_egreso': r'potasio[:\s]*(\d+(?:\.\d+)?)',
        'cloro_egreso': r'cloro[:\s]*(\d+(?:\.\d+)?)',
        'creatinina_egreso': r'creatinina[:\s]*(\d+(?:\.\d+)?)',
        'bun_egreso': r'bun[:\s]*(\d+(?:\.\d+)?)',
        'urea_egreso': r'urea[:\s]*(\d+(?:\.\d+)?)',
        'glucosa_egreso': r'glucosa[:\s]*(\d+(?:\.\d+)?)',
        'bilirrubina_total_egreso': r'bilirrubina total[:\s]*(\d+(?:\.\d+)?)',
        'albumina_egreso': r'alb[úu]mina[:\s]*(\d+(?:\.\d+)?)',
        'gasometria_ph_egreso': r'ph[:\s]*(\d+(?:\.\d+)?)',
        'gasometria_pco2_egreso': r'pco2[:\s]*(\d+(?:\.\d+)?)',
        'gasometria_po2_egreso': r'po2[:\s]*(\d+(?:\.\d+)?)',
        'gasometria_hco3_egreso': r'hco3[:\s]*(\d+(?:\.\d+)?)',
        'gasometria_lactato_egreso': r'lactato[:\s]*(\d+(?:\.\d+)?)',
    }
    for key, pat in patterns.items():
        match = re.search(pat, text, re.IGNORECASE)
        if match:
            try:
                val_str = match.group(1).strip()
                if re.match(r'^\d+(?:\.\d+)?$', val_str):
                    if '.' in val_str:
                        data[key] = float(val_str)
                    else:
                        data[key] = int(val_str)
                else:
                    data[key] = val_str
            except:
                data[key] = match.group(1).strip()
    return data

def process_discharge_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        doc = Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
    elif ext in ['.jpg', '.jpeg', '.png']:
        text = extract_text_from_image(file_path)
    else:
        return None, None, None
    with get_db_connection() as conn:
        patient_id, patient_data = identify_patient_from_text(text, conn)
    if not patient_id:
        return None, None, None
    discharge_data = extract_discharge_data(text)
    return patient_id, patient_data, discharge_data

def create_discharge_dataframe(discharge_list):
    columns = ['archivo', 'paciente', 'expediente'] + egreso_field_names + ['notas']
    df = pd.DataFrame(discharge_list)
    for col in columns:
        if col not in df.columns:
            df[col] = ""
    df = df[[c for c in columns if c in df.columns]]
    return df

def save_discharge_data(df):
    resultados = []
    for idx, row in df.iterrows():
        expediente = row.get('expediente')
        if expediente and pd.notna(expediente) and expediente != "":
            with get_db_connection() as conn:
                row_db = conn.execute("SELECT id FROM patients WHERE json_extract(data, '$.expediente') = ?", (expediente,)).fetchone()
            if row_db:
                patient_id = row_db['id']
            else:
                resultados.append(f"Fila {idx+1}: Paciente con expediente {expediente} no encontrado")
                continue
        else:
            nombre = row.get('paciente')
            if nombre and pd.notna(nombre) and nombre != "":
                with get_db_connection() as conn:
                    rows = conn.execute("SELECT id FROM patients WHERE json_extract(data, '$.nombre_completo') LIKE ?", (f'%{nombre}%',)).fetchall()
                if rows:
                    patient_id = rows[0]['id']
                else:
                    resultados.append(f"Fila {idx+1}: Paciente con nombre {nombre} no encontrado")
                    continue
            else:
                resultados.append(f"Fila {idx+1}: Sin identificador de paciente")
                continue
        current = get_patient(patient_id)
        if not current:
            resultados.append(f"Fila {idx+1}: Error al recuperar paciente")
            continue
        current_data = current["data"]
        updated = False
        for col in egreso_field_names:
            if col in row and pd.notna(row[col]) and row[col] != "":
                current_data[col] = row[col]
                updated = True
        if updated:
            update_patient(patient_id, current_data)
            resultados.append(f"Fila {idx+1}: Datos de egreso actualizados para {current_data.get('nombre_completo', '')} (ID: {patient_id})")
        else:
            resultados.append(f"Fila {idx+1}: No se encontraron datos de egreso para actualizar")
        if pd.notna(row.get('notas')) and row['notas'].strip():
            add_note(patient_id, row['notas'].strip(), note_type='egreso')
    return "\n".join(resultados)

# ==================== INTERFAZ GRADIO ====================
init_db()

with gr.Blocks(title="Procesador de Censos y Egresos UCI", theme=gr.themes.Soft()) as demo:
    gr.Markdown("# 🏥 Procesador de Censos y Notas de Egreso UCI")

    with gr.Tab("Censo UCI (ingreso/actualización)"):
        with gr.Row():
            file_input = gr.File(label="Archivos Word (.docx)", file_count="multiple", file_types=[".docx"])
            process_btn = gr.Button("🔍 Procesar archivos", variant="primary")
        with gr.Row():
            df_editor = gr.Dataframe(label="Datos de pacientes (editable)", interactive=True, wrap=True, col_count=(20, "fixed"))
        with gr.Row():
            status_output = gr.Textbox(label="Estado", interactive=False)
        with gr.Row():
            save_btn = gr.Button("💾 Guardar en base de datos", variant="secondary")
            save_output = gr.Textbox(label="Resultado del guardado", interactive=False)

        def process_files(files):
            if not files:
                return None, "No se seleccionaron archivos"
            merged_patients, total_records = process_multiple_docx(files)
            df = create_dataframe_from_patients(merged_patients)
            return df, f"Se procesaron {len(files)} archivo(s) con {total_records} registros, fusionados en {len(merged_patients)} paciente(s) únicos."

        def save_from_df(df):
            if df is None or df.empty:
                return "No hay datos para guardar"
            resultado = save_to_database_from_df(df)
            return resultado

        process_btn.click(fn=process_files, inputs=file_input, outputs=[df_editor, status_output])
        save_btn.click(fn=save_from_df, inputs=df_editor, outputs=save_output)

    with gr.Tab("Nota de Egreso"):
        with gr.Row():
            discharge_file_input = gr.File(label="Archivos de egreso (.docx, .jpg, .png)", file_count="multiple", file_types=[".docx", ".jpg", ".jpeg", ".png"])
            discharge_process_btn = gr.Button("🔍 Procesar archivos", variant="primary")
        with gr.Row():
            discharge_df_editor = gr.Dataframe(label="Datos de egreso (editable)", interactive=True, wrap=True, col_count=(20, "fixed"))
        with gr.Row():
            discharge_status = gr.Textbox(label="Estado", interactive=False)
        with gr.Row():
            discharge_save_btn = gr.Button("💾 Guardar en base de datos", variant="secondary")
            discharge_save_output = gr.Textbox(label="Resultado del guardado", interactive=False)

        def process_discharge_files(files):
            if not files:
                return None, "No se seleccionaron archivos"
            results = []
            for f in files:
                patient_id, patient_data, discharge_data = process_discharge_file(f.name)
                if patient_id:
                    discharge_data['archivo'] = os.path.basename(f.name)
                    discharge_data['paciente'] = patient_data.get('nombre_completo', '')
                    discharge_data['expediente'] = patient_data.get('expediente', '')
                    results.append(discharge_data)
                else:
                    results.append({'archivo': os.path.basename(f.name), 'error': 'No se pudo identificar al paciente'})
            df = create_discharge_dataframe(results)
            return df, f"Se procesaron {len(files)} archivo(s). Revise y edite."

        def save_discharge_from_df(df):
            if df is None or df.empty:
                return "No hay datos para guardar"
            resultado = save_discharge_data(df)
            return resultado

        discharge_process_btn.click(fn=process_discharge_files, inputs=discharge_file_input, outputs=[discharge_df_editor, discharge_status])
        discharge_save_btn.click(fn=save_discharge_from_df, inputs=discharge_df_editor, outputs=discharge_save_output)

demo.launch(debug=True, share=False)
